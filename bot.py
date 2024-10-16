import pytesseract
from pdf2image import convert_from_path
from PIL import Image, ImageEnhance, ImageFilter  # لتحسين الصور
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from bidi.algorithm import get_display
from telegram import Update, InputFile
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
import requests
import os

# إعداد مسار Tesseract (على Windows)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def enhance_image(image_path):
    """تحسين الصورة لتكون أكثر وضوحًا."""
    img = Image.open(image_path)
    img = img.convert('L')  # تحويل إلى تدرجات الرمادي
    img = ImageEnhance.Contrast(img).enhance(2)  # زيادة التباين
    img = img.filter(ImageFilter.SHARPEN)  # تحسين الحدة
    return img

def extract_text_from_image(image_path):
    """استخراج النص من صورة بعد تحسينها."""
    img = enhance_image(image_path)
    config = '--psm 6'  # إعداد Tesseract لتحسين الاستخراج
    text = pytesseract.image_to_string(img, lang='ara', config=config)
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text

def extract_text_from_pdf(pdf_path):
    """استخراج النص من PDF."""
    images = convert_from_path(pdf_path)
    all_text = ""
    for image in images:
        # حفظ الصورة بشكل مؤقت لتحسينها
        temp_image_path = "temp_image.png"
        image.save(temp_image_path)
        all_text += extract_text_from_image(temp_image_path) + "\n\n"
        os.remove(temp_image_path)
    return all_text

def save_to_word(text, output_path):
    """حفظ النص في ملف Word مع تنسيق مناسب."""
    doc = Document()
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(text)
    run.font.name = 'Arial'  # خط مناسب للنص العربي
    run.font.size = Pt(14)  # حجم الخط

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # محاذاة إلى اليمين
    set_paragraph_rtl(paragraph)

    doc.save(output_path)

def set_paragraph_rtl(paragraph):
    """ضبط اتجاه النص إلى اليمين إلى اليسار."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph_element = paragraph._element
    rtl = OxmlElement('w:bidi')
    rtl.set(qn('w:val'), '1')
    paragraph_element.append(rtl)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """الرد على أمر /start."""
    await update.message.reply_text("أهلاً بك! أرسل لي صورة أو ملف PDF لاستخراج النص وتحويله إلى وورد.")

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة الملفات الواردة."""
    try:
        file = await update.message.document.get_file() if update.message.document else await update.message.photo[-1].get_file()
        file_path = file.file_path

        local_filename = os.path.basename(file_path)
        response = requests.get(file_path)
        with open(local_filename, 'wb') as f:
            f.write(response.content)

        if local_filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(local_filename)
        else:
            extracted_text = extract_text_from_image(local_filename)

        output_word_path = "output.docx"
        save_to_word(extracted_text, output_word_path)

        with open(output_word_path, 'rb') as word_file:
            await update.message.reply_document(document=InputFile(word_file), filename="output.docx")

        os.remove(local_filename)

    except Exception as e:
        await update.message.reply_text(f"حدث خطأ أثناء معالجة الملف: {str(e)}")

def main():
    """تشغيل البوت."""
    TOKEN = "7648168005:AAGnTMOudmGjUNB2XfnZv42AvFZKi9psX2E"
    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL | filters.PHOTO, handle_file))

    app.run_polling()

if __name__ == "__main__":
    main()
